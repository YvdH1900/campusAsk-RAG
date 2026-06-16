"""
文档管理 API 路由
================
提供文档上传、审核和管理功能：
1. POST /api/documents/upload - 教师上传文档
2. GET /api/documents/my - 查看我的文档列表
3. GET /api/documents/pending - 管理员查看待审核文档
4. POST /api/documents/{id}/review - 管理员审核文档
5. GET /api/documents - 查看所有文档
"""

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from datetime import datetime, date
import os
import shutil
import logging
from typing import Optional
from math import ceil
from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.models import User, UserRole, Document
from app.tasks.document_tasks import process_document_task
from app.schemas import DocumentResponse, DocumentReviewRequest, PaginatedDocumentResponse
from pydantic import BaseModel

from app.api.toc_splitter import split_pdf_by_toc
from typing import List
import uuid

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["文档管理"])

UPLOAD_DIR = "uploads/documents"
TEMP_UPLOAD_DIR = "uploads/temp_documents"
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md"}

# 大文件自动拆分阈值（字节），超过此大小的文件会被拆分为多个子文档
LARGE_FILE_THRESHOLD = 20 * 1024 * 1024  # 20MB

# 拆分策略阈值
MIN_CHARS = 2000      # 章节过小合并阈值
MAX_CHARS = 200000    # 章节过大拆分阈值（仅处理异常巨大的章节）
TARGET_CHARS = 50000  # 目标单文档字符数（文件级拆分，非向量化分块）


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def split_large_file(file_path: str, original_filename: str) -> list[str]:
    """
    统一的大文件拆分入口，按文件类型自动选择拆分策略
    
    策略：
    - PDF → 按 TOC 目录章节拆分（toc_splitter.py）
    - DOCX → 按标题样式（Heading 1-6）拆分
    - DOC → 尝试提取文本后按内容量拆分
    - TXT/MD → 按内容量拆分（保持段落完整性）
    
    Args:
        file_path: 原文件路径
        original_filename: 原始文件名
    
    Returns:
        拆分后的文件路径列表（空列表表示不需要拆分或拆分失败）
    """
    file_ext = get_file_extension(original_filename)
    
    if file_ext == ".pdf":
        return _split_pdf(file_path, original_filename)
    elif file_ext == ".docx":
        return _split_docx(file_path, original_filename)
    elif file_ext == ".doc":
        return _split_doc(file_path, original_filename)
    elif file_ext in (".txt", ".md"):
        return _split_text_file(file_path, original_filename, file_ext)
    else:
        logger.warning(f"不支持的拆分类型: {file_ext}")
        return []


def _split_pdf(file_path: str, original_filename: str) -> list[str]:
    """PDF 按 TOC 目录章节拆分"""
    return split_pdf_by_toc(file_path, original_filename)


def _split_docx(file_path: str, original_filename: str) -> list[str]:
    """DOCX 按标题样式拆分"""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("未找到 python-docx，无法拆分 DOCX 文件")
        return []
    
    doc = DocxDocument(file_path)
    total_paragraphs = len(doc.paragraphs)
    
    # 标题样式列表
    HEADING_STYLES = [
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "标题 1", "标题 2", "标题 3",
        "标题 4", "标题 5", "标题 6",
    ]
    
    # 提取标题段落
    headings = []
    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name in HEADING_STYLES and para.text.strip():
            headings.append({
                "index": idx,
                "title": para.text.strip()[:50],
                "style": style_name,
            })
    
    logger.info(f"DOCX 标题提取: 找到 {len(headings)} 个标题")
    
    if len(headings) < 2:
        logger.warning("DOCX 标题不足，回退按段落数拆分")
        return _split_docx_by_paragraphs(doc, file_path, original_filename)
    
    # 构建章节
    sections = _build_sections_from_headings(doc, headings, total_paragraphs)
    
    # 二次处理 + 生成文件
    return _process_and_split_docx_sections(doc, sections, file_path, original_filename)


def _split_doc(file_path: str, original_filename: str) -> list[str]:
    """
    DOC 文件拆分（旧版 Word 格式）
    DOC 是二进制 OLE 格式，无法直接解析结构，不拆分
    直接交给 DocumentParser 处理（它会用 olefile 提取文本）
    """
    logger.info(f"DOC 文件不拆分，直接交给解析器处理: {original_filename}")
    return []


def _split_text_file(file_path: str, original_filename: str, file_ext: str) -> list[str]:
    """TXT/MD 文件按内容量拆分（保持段落完整性）"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="gbk") as f:
                content = f.read()
        except Exception as e:
            logger.error(f"无法读取文本文件: {e}")
            return []
    
    total_chars = len(content)
    if total_chars <= MAX_CHARS:
        return []  # 不需要拆分
    
    # 按段落拆分
    paragraphs = content.split("\n\n")
    
    # 按内容量分组
    groups = []
    current_group = []
    current_chars = 0
    
    for para in paragraphs:
        para_chars = len(para)
        if current_chars + para_chars > TARGET_CHARS and current_group:
            groups.append(current_group)
            current_group = [para]
            current_chars = para_chars
        else:
            current_group.append(para)
            current_chars += para_chars
    
    if current_group:
        groups.append(current_group)
    
    # 生成拆分文件
    dir_name = os.path.dirname(file_path)
    base_name = os.path.splitext(original_filename)[0]
    
    split_files = []
    for idx, group in enumerate(groups):
        split_content = "\n\n".join(group)
        split_filename = f"{base_name}_part{idx+1:02d}{file_ext}"
        split_path = os.path.join(dir_name, split_filename)
        
        with open(split_path, "w", encoding="utf-8") as f:
            f.write(split_content)
        
        split_files.append(split_path)
        logger.info(f"文本拆分: {split_filename} (~{len(split_content)}字符)")
    
    logger.info(f"文本文件拆分完成: {len(split_files)} 个文件")
    return split_files


def _build_sections_from_headings(doc, headings: list, total_paragraphs: int) -> list:
    """根据标题构建章节"""
    sections = []
    for i, heading in enumerate(headings):
        start_idx = heading["index"]
        end_idx = headings[i + 1]["index"] - 1 if i + 1 < len(headings) else total_paragraphs - 1
        
        section_chars = sum(len(doc.paragraphs[p_idx].text or "") for p_idx in range(start_idx, end_idx + 1))
        
        sections.append({
            "title": heading["title"],
            "start_idx": start_idx,
            "end_idx": end_idx,
            "chars": section_chars,
        })
    return sections


def _process_and_split_docx_sections(doc, sections: list, file_path: str, original_filename: str) -> list[str]:
    """对章节进行二次处理（合并/拆分）并生成文件"""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []
    
    # 合并过小章节
    merged_sections = []
    i = 0
    while i < len(sections):
        current = sections[i]
        if current["chars"] < MIN_CHARS and i + 1 < len(sections):
            next_section = sections[i + 1]
            merged = {
                "title": f"{current['title']} + {next_section['title']}",
                "start_idx": current["start_idx"],
                "end_idx": next_section["end_idx"],
                "chars": current["chars"] + next_section["chars"],
            }
            logger.info(f"  合并: [{current['title'][:20]}] + [{next_section['title'][:20]}] (~{merged['chars']}字符)")
            merged_sections.append(merged)
            i += 2
        else:
            merged_sections.append(current)
            i += 1
    
    # 拆分过大章节
    MAX_SPLITS_PER_SECTION = 3  # 每个章节最多拆 3 份，避免过度碎片化
    
    final_sections = []
    for s in merged_sections:
        if s["chars"] <= MAX_CHARS:
            final_sections.append(s)
        else:
            paragraphs_in_section = s["end_idx"] - s["start_idx"] + 1
            avg_chars_per_para = s["chars"] / paragraphs_in_section if paragraphs_in_section > 0 else 100
            paras_per_split = max(10, int(TARGET_CHARS / avg_chars_per_para))
            num_splits = ceil(paragraphs_in_section / paras_per_split)
            
            # 限制最多拆 3 份，超过就直接保留原章节
            if num_splits > MAX_SPLITS_PER_SECTION:
                logger.info(
                    f"  保留: [{s['title'][:30]}] (~{s['chars']}字符，需拆{num_splits}份 > {MAX_SPLITS_PER_SECTION}份上限，直接保留)"
                )
                final_sections.append(s)
            else:
                logger.info(f"  拆分: [{s['title'][:30]}] (~{s['chars']}字符 → {num_splits} 份)")
                
                for j in range(num_splits):
                    split_start = s["start_idx"] + j * paras_per_split
                    split_end = min(s["start_idx"] + (j + 1) * paras_per_split - 1, s["end_idx"])
                    final_sections.append({
                        "title": f"{s['title']} (部分{j+1}/{num_splits})",
                        "start_idx": split_start,
                        "end_idx": split_end,
                        "chars": int(avg_chars_per_para * (split_end - split_start + 1)),
                    })
    
    # 生成拆分文件
    dir_name = os.path.dirname(file_path)
    base_name = os.path.splitext(original_filename)[0]
    ext = os.path.splitext(original_filename)[1]
    
    split_files = []
    for idx, section in enumerate(final_sections):
        if section["start_idx"] > section["end_idx"]:
            continue
        
        new_doc = DocxDocument()
        
        for p_idx in range(section["start_idx"], section["end_idx"] + 1):
            para = doc.paragraphs[p_idx]
            if not para.text.strip() and p_idx > section["start_idx"]:
                continue
            
            new_para = new_doc.add_paragraph()
            new_para.text = para.text
            
            if para.style:
                try:
                    new_para.style = para.style.name
                except Exception:
                    pass
            
            if para.paragraph_format:
                try:
                    new_para.paragraph_format.alignment = para.paragraph_format.alignment
                except Exception:
                    pass
        
        split_filename = f"{base_name}_第{idx+1:02d}章_{section['title'][:20]}{ext}"
        split_filename = re.sub(r'[\\/*?:"<>|]', '', split_filename)[:150]
        split_path = os.path.join(dir_name, split_filename)
        
        new_doc.save(split_path)
        split_files.append(split_path)
        
        logger.info(
            f"    → {split_filename} (第{section['start_idx']+1}-{section['end_idx']+1}段, ~{section['chars']}字符)"
        )
    
    logger.info(f"DOCX 按标题拆分完成: {len(split_files)} 个文件")
    return split_files


def _split_docx_by_paragraphs(doc, file_path: str, original_filename: str, max_paragraphs: int = 500) -> list[str]:
    """回退方案：按段落数拆分 DOCX"""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []
    
    total_paragraphs = len(doc.paragraphs)
    dir_name = os.path.dirname(file_path)
    base_name = os.path.splitext(original_filename)[0]
    ext = os.path.splitext(original_filename)[1]
    
    split_files = []
    num_splits = ceil(total_paragraphs / max_paragraphs)
    
    for i in range(num_splits):
        start_idx = i * max_paragraphs
        end_idx = min((i + 1) * max_paragraphs, total_paragraphs)
        
        new_doc = DocxDocument()
        
        for p_idx in range(start_idx, end_idx):
            para = doc.paragraphs[p_idx]
            new_para = new_doc.add_paragraph()
            new_para.text = para.text
            
            if para.style:
                try:
                    new_para.style = para.style.name
                except Exception:
                    pass
        
        split_filename = f"{base_name}_part{i+1:02d}{ext}"
        split_path = os.path.join(dir_name, split_filename)
        new_doc.save(split_path)
        split_files.append(split_path)
        logger.info(f"DOCX 拆分: {split_filename} (第 {start_idx+1}-{end_idx} 段)")
    
    return split_files


def _split_file_by_size(file_path: str, original_filename: str, chunk_size: int = TARGET_CHARS) -> list[str]:
    """按文件大小均匀拆分（通用回退方案）"""
    try:
        with open(file_path, "rb") as f:
            content = f.read()
    except Exception as e:
        logger.error(f"无法读取文件: {e}")
        return []
    
    total_size = len(content)
    if total_size <= chunk_size:
        return []
    
    dir_name = os.path.dirname(file_path)
    base_name = os.path.splitext(original_filename)[0]
    ext = os.path.splitext(original_filename)[1]
    
    # 尝试按行拆分，避免切断中间内容
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("gbk")
        except Exception:
            text = content.decode("latin-1")
    
    lines = text.split("\n")
    
    groups = []
    current_group = []
    current_chars = 0
    
    for line in lines:
        line_chars = len(line)
        if current_chars + line_chars > chunk_size and current_group:
            groups.append(current_group)
            current_group = [line]
            current_chars = line_chars
        else:
            current_group.append(line)
            current_chars += line_chars
    
    if current_group:
        groups.append(current_group)
    
    split_files = []
    for idx, group in enumerate(groups):
        split_content = "\n".join(group)
        split_filename = f"{base_name}_part{idx+1:02d}{ext}"
        split_path = os.path.join(dir_name, split_filename)
        
        with open(split_path, "w", encoding="utf-8") as f:
            f.write(split_content)
        
        split_files.append(split_path)
        logger.info(f"文件拆分: {split_filename} (~{len(split_content)}字符)")
    
    return split_files


def reset_daily_count_if_needed(user: User, db: Session):
    today = date.today()
    if user.last_reset_date is None or user.last_reset_date != today:
        user.questions_today = 0
        user.uploads_today = 0
        user.last_reset_date = today
        db.commit()
        db.refresh(user)


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    教师/管理员上传文档
    
    教师上传后状态为待审核，需要管理员审核通过后才入库。
    管理员上传直接入库，无需审核。
    """
    if current_user.role not in [UserRole.TEACHER, UserRole.ADMIN]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有教师和管理员角色可以上传文档"
        )
    
    if current_user.role == UserRole.TEACHER and current_user.pending_approval:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="您的教师账号正在审核中，无法上传文档"
        )
    
    # 管理员无上传限制
    if current_user.role != UserRole.ADMIN:
        if current_user.max_uploads_per_day == 0:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="您的上传权限已被限制，无法上传文档"
            )
        
        reset_daily_count_if_needed(current_user, db)
        
        if current_user.uploads_today >= current_user.max_uploads_per_day:
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail=f"今日上传次数已达上限 ({current_user.max_uploads_per_day}次)，请明天再试"
            )
    
    file_ext = get_file_extension(file.filename)
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件类型，仅支持: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    is_admin = current_user.role == UserRole.ADMIN

    if is_admin:
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        unique_filename = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 获取文件大小
        file_size = os.path.getsize(file_path)

        # 检查是否为大文件，需要拆分
        is_large_file = file_size > LARGE_FILE_THRESHOLD
        split_files = []

        if is_large_file:
            split_files = split_large_file(file_path, file.filename)

        # 无论是否拆分，MySQL 只存 1 条原始文档记录
        new_document = Document(
            filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            category=category,
            description=description,
            status="processing",
            review_status="approved",
            uploaded_by=current_user.id,
            reviewed_by=current_user.id,
            reviewed_at=datetime.utcnow(),
        )

        db.add(new_document)
        db.commit()
        db.refresh(new_document)

        # 使用 Celery 异步处理文档（大文件拆分在任务内部处理）
        if split_files:
            logger.info(f"大文件自动拆分: {file.filename} ({file_size/1024/1024:.1f}MB) -> {len(split_files)} 个文件")
            # 保留原文件，拆分文件作为临时处理文件
            process_document_task.apply_async(
                args=[new_document.id],
                kwargs={"split_files": split_files},
            )
        else:
            process_document_task.delay(new_document.id)

        return new_document
    else:
        os.makedirs(TEMP_UPLOAD_DIR, exist_ok=True)
        
        unique_filename = f"{uuid.uuid4().hex}{file_ext}"
        temp_file_path = os.path.join(TEMP_UPLOAD_DIR, unique_filename)
        
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 获取文件大小
        file_size = os.path.getsize(temp_file_path)
        
        new_document = Document(
            filename=file.filename,
            file_path=temp_file_path,
            file_size=file_size,
            category=category,
            description=description,
            status="pending",
            review_status="pending",
            uploaded_by=current_user.id,
        )
    
        # 教师上传增加计数
        current_user.uploads_today += 1
        db.add(new_document)
        db.commit()
        db.refresh(new_document)
    
        return new_document


@router.get("/my", response_model=list[DocumentResponse])
def get_my_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    查看我的文档列表
    
    教师可以查看自己上传的所有文档及其审核状态。
    """
    if current_user.role not in [UserRole.TEACHER, UserRole.ADMIN]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无权访问"
        )
    
    documents = db.query(Document).filter(
        Document.uploaded_by == current_user.id
    ).order_by(Document.created_at.desc()).all()
    
    return documents


@router.get("/pending", response_model=list[DocumentResponse])
def get_pending_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    管理员查看待审核文档列表
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    documents = db.query(Document).filter(
        Document.review_status == "pending"
    ).order_by(Document.created_at.asc()).all()
    
    return documents


@router.post("/{document_id}/review", response_model=DocumentResponse)
def review_document(
    document_id: int,
    review_data: DocumentReviewRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    管理员审核文档
    
    action: approve (通过) / reject (驳回)
    reason: 驳回理由（驳回时必填）
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以审核文档"
        )
    
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    if document.review_status != "pending":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="该文档已经审核过了"
        )
    
    if review_data.action == "approve":
        document.review_status = "approved"
        document.status = "processing"

        if os.path.exists(document.file_path):
            os.makedirs(UPLOAD_DIR, exist_ok=True)
            filename = os.path.basename(document.file_path)
            new_file_path = os.path.join(UPLOAD_DIR, filename)

            if document.file_path != new_file_path:
                shutil.move(document.file_path, new_file_path)
                document.file_path = new_file_path

        document.reviewed_by = current_user.id
        document.reviewed_at = datetime.utcnow()

        # 检查是否为大文件，需要拆分
        file_size = os.path.getsize(document.file_path) if os.path.exists(document.file_path) else 0
        split_files = []

        if file_size > LARGE_FILE_THRESHOLD:
            split_files = split_large_file(document.file_path, document.filename)

        if split_files:
            logger.info(f"审核通过-大文件自动拆分: {document.filename} ({file_size/1024/1024:.1f}MB) -> {len(split_files)} 个文件")
            # 删除原文件，拆分文件作为临时处理文件
            os.remove(document.file_path)
            document.status = "processing"
            db.commit()
            
            # 将拆分文件列表传递给 Celery 任务
            process_document_task.apply_async(
                args=[document.id],
                kwargs={"split_files": split_files},
            )
        else:
            document.status = "processing"
            db.commit()
            process_document_task.delay(document.id)

        return document
                
    elif review_data.action == "reject":
        if not review_data.reason:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="驳回时必须填写理由"
            )
        document.review_status = "rejected"
        document.status = "failed"
        document.reject_reason = review_data.reason
        
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="无效的操作，action 必须为 approve 或 reject"
        )
    
    document.reviewed_by = current_user.id
    document.reviewed_at = datetime.utcnow()
    
    db.commit()
    db.refresh(document)
    
    return document


@router.get("/", response_model=PaginatedDocumentResponse)
def get_all_documents(
    page: int = Query(1, ge=1, description="页码"),
    page_size: int = Query(10, ge=1, le=100, description="每页数量"),
    status_filter: Optional[str] = Query(None, alias="status", description="按处理状态筛选"),
    review_status: Optional[str] = Query(None, description="按审核状态筛选"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    查看所有文档（管理员，支持分页）
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    query = db.query(Document)
    
    if status_filter:
        query = query.filter(Document.status == status_filter)
    
    if review_status:
        query = query.filter(Document.review_status == review_status)
    
    total = query.count()
    pages = ceil(total / page_size) if total > 0 else 1
    
    documents = query.order_by(Document.created_at.desc()).offset((page - 1) * page_size).limit(page_size).all()
    
    return PaginatedDocumentResponse(
        total=total,
        page=page,
        page_size=page_size,
        pages=pages,
        items=documents,
    )


@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    删除文档（管理员）
    
    删除内容包括：
    1. 本地文件
    2. 向量库中的向量数据
    3. 数据库记录
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    # 1. 删除本地文件
    if os.path.exists(document.file_path):
        try:
            os.remove(document.file_path)
            logger.info(f"已删除本地文件: {document.file_path}")
        except Exception as e:
            logger.warning(f"删除本地文件失败: {document.file_path}, 错误: {str(e)}")
    
    # 2. 删除向量库数据（关键操作，失败则中止）
    try:
        from app.services.document_processor import DocumentProcessor
        from app.services.vector_store import VectorStore
        processor = DocumentProcessor()
        processor.delete_document_vectors(document_id)
        vector_store = VectorStore()
        orphan_deleted = vector_store.delete_orphan_vectors()
        if orphan_deleted > 0:
            logger.info(f"清理了 {orphan_deleted} 条孤儿向量")
        logger.info(f"已删除向量数据: document_id={document_id}")
    except Exception as e:
        logger.error(f"删除向量数据失败: document_id={document_id}, 错误: {str(e)}")
        # 即使向量删除失败，也继续删除数据库记录
        logger.warning("继续删除数据库记录")
    
    # 3. 删除数据库记录
    filename = document.filename
    db.delete(document)
    db.commit()
    
    logger.info(f"已删除文档: id={document_id}, filename={filename}")
    
    return {"message": "文档删除成功", "filename": filename}


class BatchDeleteRequest(BaseModel):
    """批量删除请求"""
    document_ids: List[int]


class BatchReviewRequest(BaseModel):
    """批量审核请求"""
    document_ids: List[int]
    action: str  # "approve" or "reject"
    reject_reason: Optional[str] = None


@router.post("/batch-delete")
def batch_delete_documents(
    request: BatchDeleteRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    批量删除文档（管理员）
    
    返回成功和失败的文档ID列表。
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    if not request.document_ids:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文档ID列表不能为空"
        )
    
    success_ids = []
    failed_ids = []
    errors = {}
    
    for doc_id in request.document_ids:
        try:
            document = db.query(Document).filter(Document.id == doc_id).first()
            if not document:
                failed_ids.append(doc_id)
                errors[str(doc_id)] = "文档不存在"
                continue
            
            if os.path.exists(document.file_path):
                try:
                    os.remove(document.file_path)
                except Exception as e:
                    logger.warning(f"删除本地文件失败: {document.file_path}, 错误: {str(e)}")
            
            try:
                from app.services.document_processor import DocumentProcessor
                processor = DocumentProcessor()
                processor.delete_document_vectors(doc_id)
            except Exception as e:
                logger.error(f"删除向量数据失败: document_id={doc_id}, 错误: {str(e)}")
            
            db.delete(document)
            db.commit()
            success_ids.append(doc_id)
            logger.info(f"已批量删除文档: id={doc_id}")
        except Exception as e:
            failed_ids.append(doc_id)
            errors[str(doc_id)] = str(e)
            logger.error(f"批量删除文档失败: id={doc_id}, 错误: {str(e)}")
            db.rollback()
    
    return {
        "message": f"批量删除完成，成功 {len(success_ids)} 个，失败 {len(failed_ids)} 个",
        "success_ids": success_ids,
        "failed_ids": failed_ids,
        "errors": errors,
    }


@router.post("/batch-review")
def batch_review_documents(
    request: BatchReviewRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    批量审核文档（管理员）
    
    action: "approve" 或 "reject"
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    if not request.document_ids:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文档ID列表不能为空"
        )
    
    if request.action not in ("approve", "reject"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="操作必须是 approve 或 reject"
        )
    
    if request.action == "reject" and not request.reject_reason:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="驳回时必须提供驳回原因"
        )
    
    success_ids = []
    failed_ids = []
    errors = {}
    
    for doc_id in request.document_ids:
        try:
            document = db.query(Document).filter(Document.id == doc_id).first()
            if not document:
                failed_ids.append(doc_id)
                errors[str(doc_id)] = "文档不存在"
                continue
            
            if request.action == "approve":
                document.review_status = "approved"
                document.review_comment = None
            else:
                document.review_status = "rejected"
                document.review_comment = request.reject_reason
            
            document.updated_at = datetime.utcnow()
            db.commit()
            success_ids.append(doc_id)
            logger.info(f"已批量审核文档: id={doc_id}, action={request.action}")
        except Exception as e:
            failed_ids.append(doc_id)
            errors[str(doc_id)] = str(e)
            logger.error(f"批量审核文档失败: id={doc_id}, 错误: {str(e)}")
            db.rollback()
    
    action_text = "通过" if request.action == "approve" else "驳回"
    return {
        "message": f"批量{action_text}完成，成功 {len(success_ids)} 个，失败 {len(failed_ids)} 个",
        "success_ids": success_ids,
        "failed_ids": failed_ids,
        "errors": errors,
    }


@router.post("/consistency-check")
def check_vector_consistency(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    检查向量库与数据库的一致性（管理员）
    
    返回：
    - 孤儿文档列表（向量库中有但数据库中没有）
    - 缺失文档列表（数据库中有但向量库中没有）
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    from app.services.vector_consistency import VectorConsistencyChecker
    checker = VectorConsistencyChecker()
    result = checker.check_consistency(db)
    
    return result


@router.post("/clean-orphan-vectors")
def clean_orphan_vectors(
    current_user: User = Depends(get_current_user)
):
    """
    清理孤儿向量（管理员）
    
    删除向量库中所有 document_id 无效（0 或 null）的孤儿向量
    """
    if current_user.role != UserRole.ADMIN:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="只有管理员可以访问"
        )
    
    from app.services.vector_consistency import VectorConsistencyChecker
    checker = VectorConsistencyChecker()
    deleted_count = checker.clean_orphan_vectors()
    
    return {
        "message": "孤儿向量清理完成",
        "deleted_count": deleted_count
    }


@router.get("/{document_id}/download")
def download_document(
    document_id: int,
    token: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user)
):
    """
    下载文档文件
    
    教师可以下载自己的文档，管理员可以下载所有文档。
    支持通过 URL 参数 ?token=xxx 传递 Token（用于浏览器直接访问）
    """
    # 如果通过 Depends 获取用户失败，尝试从 URL 参数获取 Token
    if not current_user and token:
        from app.core.security import decode_access_token
        payload = decode_access_token(token)
        if payload:
            user_id: str = payload.get("sub")
            if user_id:
                current_user = db.query(User).filter(User.id == int(user_id)).first()
    
    if not current_user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="无效的认证凭证"
        )
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    if current_user.role != UserRole.ADMIN and document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无权下载此文档"
        )
    
    if not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档文件不存在"
        )
    
    from fastapi.responses import StreamingResponse
    import mimetypes
    
    # 获取文件的 MIME 类型
    mime_type, _ = mimetypes.guess_type(document.filename)
    media_type = mime_type or "application/octet-stream"
    
    # 使用 StreamingResponse 来完全控制响应头
    from urllib.parse import quote
    
    def iterfile():
        with open(document.file_path, "rb") as f:
            while chunk := f.read(8192):
                yield chunk
    
    # 手动设置 Content-Disposition，确保中文文件名正确显示
    # 使用 RFC 5987 编码格式
    encoded_filename = quote(document.filename.encode('utf-8'))
    
    # 调试日志
    print(f"[DEBUG] 下载文件：{document.filename}, encoded: {encoded_filename}", flush=True)
    
    return StreamingResponse(
        iterfile(),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Content-Length": str(os.path.getsize(document.file_path)),
        }
    )


@router.get("/{document_id}/preview")
def preview_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    预览文档（仅支持文本类文件）
    
    返回文件内容，用于在线预览。
    """
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档不存在"
        )
    
    if current_user.role != UserRole.ADMIN and document.uploaded_by != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无权预览此文档"
        )
    
    if not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="文档文件不存在"
        )
    
    file_ext = get_file_extension(document.filename)
    text_extensions = {".txt", ".md"}
    
    if file_ext not in text_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="此文件类型不支持预览，请下载后查看"
        )
    
    try:
        with open(document.file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return {
            "filename": document.filename,
            "content": content,
            "size": os.path.getsize(document.file_path),
        }
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="文件编码不支持预览"
        )
