"""
文档解析服务（重构版）
======================
支持多种文档格式的智能解析：

PDF:
  - 可编辑 PDF：主选 pdfplumber，备选 PyMuPDF (fitz)
  - 扫描版 PDF：PaddleOCR + PP-StructureV2（版面分析）
  - 混合页面：正文用原生提取 + 图片区域单独 OCR
  - 低质量页面（模糊、识别率 < 60%）直接拦截

DOCX/DOC:
  - 主选 python-docx
  - 正文用原生提取，段落中的嵌入图片单独 OCR

禁止：Excel/CSV 格式（直接拒绝）

流程：
  1. 文件校验
  2. 预处理（剔除封面/目录/页眉页脚等）
  3. PDF 类型检测（可编辑/扫描/混合）
  4. 分页解析 + OCR 融合
  5. 粗洗
"""

import os
import re
import io
import logging
from typing import Optional, Generator, Dict, List, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

# python-magic 跨平台兼容
try:
    import magic
    _HAS_MAGIC = True
except (ImportError, OSError):
    _HAS_MAGIC = False
    magic = None

try:
    from langdetect import detect, LangDetectException
    _HAS_LANGDETECT = True
except ImportError:
    _HAS_LANGDETECT = False

from app.services.document_preprocessor import document_preprocessor
from app.services.document_cleaner import document_cleaner


class DocumentParser:
    """智能文档解析器"""

    # 支持的文件扩展名（禁止 Excel/CSV）
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
    BANNED_EXTENSIONS = {".xlsx", ".xls", ".csv", ".xlsm", ".xltx", ".xlt"}

    # 最大文件大小（100MB）
    MAX_FILE_SIZE = 100 * 1024 * 1024

    # 大文件拆分阈值（50MB）
    LARGE_FILE_THRESHOLD = 50 * 1024 * 1024

    # OCR 配置
    OCR_MIN_CONFIDENCE = 0.6  # 最低识别率阈值
    OCR_TEXT_MIN_LENGTH = 30  # 单页最少有效字符（低于此值视为空白/低质量）

    MIME_TYPE_MAP = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/plain",
    }

    # ==================== 文件校验 ====================

    @classmethod
    def validate_file(cls, file_path: str) -> dict:
        """
        校验文件类型和大小

        Returns:
            {"valid": bool, "mime_type": str, "size": int, "error": str}
        """
        if not os.path.exists(file_path):
            return {"valid": False, "error": "文件不存在"}

        file_size = os.path.getsize(file_path)
        if file_size > cls.MAX_FILE_SIZE:
            return {"valid": False, "error": f"文件大小超过限制 ({cls.MAX_FILE_SIZE / 1024 / 1024}MB)"}

        ext = os.path.splitext(file_path)[1].lower()

        # 禁止 Excel/CSV
        if ext in cls.BANNED_EXTENSIONS:
            return {"valid": False, "error": f"不支持 Excel/CSV 格式: {ext}"}

        if ext not in cls.ALLOWED_EXTENSIONS:
            return {"valid": False, "error": f"不支持的文件格式: {ext}"}

        if _HAS_MAGIC and magic:
            mime = magic.from_file(file_path, mime=True)
            expected_mime = cls.MIME_TYPE_MAP.get(ext)
            if expected_mime and mime != expected_mime:
                if not (expected_mime == "text/plain" and mime.startswith("text/")):
                    return {"valid": False, "error": f"文件类型不匹配，期望 {expected_mime}，实际 {mime}"}
        else:
            mime = cls.MIME_TYPE_MAP.get(ext, "application/octet-stream")
            logger.debug("python-magic 不可用，使用扩展名进行文件类型校验")

        return {"valid": True, "mime_type": mime, "size": file_size}

    # ==================== 语言检测 ====================

    @staticmethod
    def detect_language(text: str) -> str:
        if not text or len(text.strip()) < 10:
            return "unknown"
        if not _HAS_LANGDETECT:
            return "unknown"
        try:
            sample = text[:1000]
            return detect(sample)
        except Exception:
            return "unknown"

    # ==================== 主解析入口 ====================

    @staticmethod
    def parse(
        file_path: str,
        preprocess: bool = True,
        apply_coarse_clean: bool = True,
    ) -> str:
        """
        解析文档并返回纯文本内容

        Args:
            file_path: 文档文件路径
            preprocess: 是否执行预处理（剔除封面/目录等）
            apply_coarse_clean: 是否执行粗洗

        Returns:
            提取的纯文本内容
        """
        validation = DocumentParser.validate_file(file_path)
        if not validation["valid"]:
            raise ValueError(f"文件校验失败: {validation['error']}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            text = DocumentParser._parse_pdf_smart(file_path, preprocess=preprocess)
        elif ext in [".docx", ".doc"]:
            text = DocumentParser._parse_docx(file_path, preprocess=preprocess)
        elif ext == ".txt":
            text = DocumentParser._parse_txt(file_path)
        elif ext == ".md":
            text = DocumentParser._parse_md(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 粗洗
        if apply_coarse_clean and text:
            text = document_cleaner.coarse_clean(text)

        return text or ""

    # ==================== PDF 智能解析 ====================

    @staticmethod
    def _parse_pdf_smart(
        file_path: str,
        preprocess: bool = True,
    ) -> str:
        """
        PDF 智能解析：
        统一走逐页自适应路由（_parse_mixed_pdf），不再前置一刀切判类型。
        每页独立决策：有嵌入式文本且质量高 → 直接提取，否则 → OCR 兜底。
        fitz document 只打开一次，避免重复加载大文件。
        """
        fitz_doc = None
        try:
            import fitz

            # 预处理：获取有效页面索引
            valid_pages = []
            if preprocess:
                valid_pages = document_preprocessor.preprocess_pdf_pages(file_path)
            if not valid_pages:
                # 预处理排除了全部页面（常见于扫描件）
                logger.warning(f"PDF 预处理后无有效页面: {file_path}, 使用全部页面")
                fitz_doc = fitz.open(file_path)
                valid_pages = list(range(len(fitz_doc)))
            else:
                # 预处理通过了部分页面，打开 fitz 准备解析
                fitz_doc = fitz.open(file_path)

            if not valid_pages:
                return ""

            # 统一走逐页自适应路由（内部会快速重分类纯扫描件走全量 OCR 以加速）
            return DocumentParser._parse_mixed_pdf(file_path, valid_pages, fitz_doc=fitz_doc)

        except ImportError:
            # fitz 不可用，回退
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    valid_pages = list(range(len(pdf.pages)))
                    return DocumentParser._parse_mixed_pdf(file_path, valid_pages)
            except Exception:
                return ""
        except Exception as e:
            logger.error(f"PDF 智能解析失败: {file_path}, {str(e)}")
            return ""
        finally:
            if fitz_doc is not None:
                try:
                    fitz_doc.close()
                except Exception:
                    pass

    @staticmethod
    def _text_quality_score(text: str) -> float:
        """
        计算文本质量评分（0.0 ~ 1.0）

        高分 = 有意义的自然语言文本（中文为主）
        低分 = 纯数字、日期、乱码、孤立字符

        用于判断从 PDF 提取的文本是否为有效内容，
        防止扫描件 PDF 的嵌入垃圾文本层被误判为"可编辑"。
        """
        if not text or len(text) < 10:
            return 0.0

        total = len(text)

        # 中文字符（含中文标点）
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', text))
        # 常见中文标点（句号、逗号、顿号、分号、冒号、问号、感叹号、引号、书名号等）
        chinese_punct = len(re.findall(r'[，。！？；：、""''（）《》【】—…]', text))
        # 英文/数字
        digits = len(re.findall(r'\d', text))
        ascii_letters = len(re.findall(r'[a-zA-Z]', text))
        # 空白字符
        whitespace = len(re.findall(r'\s', text))
        # 孤立数字（如页码、编号）
        isolated_numbers = len(re.findall(r'(?:^|\s)\d+(?:\s|$)', text))

        # 纯数字+空白占比过高 → 垃圾文本
        junk_ratio = (digits + whitespace) / max(total, 1)
        if junk_ratio > 0.7:
            return 0.05

        # 单字符行过多 → PDF 文本层排版碎片（如 "明\n德\n厚\n学"），但若中文字符总量足够
        # 且中文占比高，说明是有效文本只是排版导致逐字符换行，不应直接判为垃圾
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            avg_chars_per_line = sum(len(l) for l in lines) / len(lines)
            if avg_chars_per_line < 2.0 and chinese_chars > 0:
                if chinese_chars > 100 and (chinese_chars / max(total, 1)) > 0.30:
                    # 中文内容充足，仅排版碎片化，给合理评分
                    return 0.25
                return 0.08

        # 中文字符占比（核心信号）
        chinese_ratio = chinese_chars / max(total, 1)

        # 中文标点占比（有标点说明是完整句子）
        punct_ratio = chinese_punct / max(total, 1)

        # 孤立数字多 → 页码/编号，不是正文
        isolated_penalty = min(0.5, isolated_numbers * 0.05)

        # 综合评分：中文占比权重 0.6，标点权重 0.25，英文权重 0.15
        english_ratio = ascii_letters / max(total, 1)
        score = chinese_ratio * 0.6 + punct_ratio * 0.25 + english_ratio * 0.15
        score = max(0.0, score - isolated_penalty)

        return min(1.0, score)

    @staticmethod
    def _is_text_likely_image(img_bytes: bytes) -> bool:
        """
        预检图片是否可能包含文字内容

        跳过低对比度图片（照片、背景图等），避免无效 OCR 调用。
        判断依据：灰度标准差（文字图片有高对比度边缘）。

        Returns:
            True 如果图片可能包含文字，值得 OCR
        """
        try:
            import numpy as np
            from PIL import Image
            import io

            pil_img = Image.open(io.BytesIO(img_bytes))
            # 只分析前 500x500 区域以加速
            if pil_img.width > 500 or pil_img.height > 500:
                pil_img = pil_img.resize((min(pil_img.width, 500), min(pil_img.height, 500)))

            gray = np.array(pil_img.convert('L'))
            std = float(gray.std())

            # 标准差 < 30 → 几乎无对比度，不可能是文字图片
            # 华中科技大学那个嵌入 JPEG 标准差约 68，但它返回 [None] 是因为
            # 文字区域太少/太模糊，仍需 OCR 判断
            if std < 30:
                logger.debug(f"跳过低对比度图片 (std={std:.1f})")
                return False

            return True
        except Exception:
            return True  # 解析失败则放行，交给 OCR 判断

    @staticmethod
    def _preprocess_ocr_image(img_bytes: bytes) -> bytes:
        """
        OCR 图像预处理：对比度增强 + 自适应二值化

        针对扫描件褪色、文字淡的问题，先增强对比度再做局部自适应二值化，
        使文字与背景分离，大幅提升 PaddleOCR 识别率。
        """
        try:
            import numpy as np
            from PIL import Image, ImageEnhance, ImageFilter
            import io

            pil_img = Image.open(io.BytesIO(img_bytes))
            original_mode = pil_img.mode

            # 转灰度
            if pil_img.mode != 'L':
                gray = pil_img.convert('L')
            else:
                gray = pil_img

            # 1. 轻度锐化（让文字边缘更清晰）
            gray = gray.filter(ImageFilter.UnsharpMask(radius=1.5, percent=80, threshold=3))

            # 2. 对比度增强
            enhancer = ImageEnhance.Contrast(gray)
            gray = enhancer.enhance(2.0)

            # 3. 自适应二值化（用高斯模糊构建局部阈值掩膜）
            arr = np.array(gray, dtype=np.float32)
            # 用大核高斯模糊作为局部均值背景
            from scipy import ndimage
            blurred = ndimage.gaussian_filter(arr, sigma=15)
            # 二值化：局部亮度 > 背景均值 → 白(255)，否则黑(0)
            # 偏移量 -5 使偏白的文字区域更容易被保留
            binary = np.where(arr > blurred - 5, 255, 0).astype(np.uint8)
            # 反色：如果大部分像素接近255（原图偏白），说明白底黑字，不用反色
            white_ratio = (binary == 255).sum() / binary.size
            if white_ratio > 0.85:
                binary = 255 - binary

            result = Image.fromarray(binary, mode='L')

            # 转回 RGB (PaddleOCR 期望 3 通道)
            result_rgb = result.convert('RGB')

            buf = io.BytesIO()
            result_rgb.save(buf, format='PNG')
            return buf.getvalue()

        except ImportError:
            # scipy 不可用时的降级方案：仅做对比度增强
            try:
                from PIL import Image, ImageEnhance
                import io

                pil_img = Image.open(io.BytesIO(img_bytes))
                if pil_img.mode != 'L':
                    gray = pil_img.convert('L')
                else:
                    gray = pil_img
                enhancer = ImageEnhance.Contrast(gray)
                gray = enhancer.enhance(2.5)
                result_rgb = gray.convert('RGB')
                buf = io.BytesIO()
                result_rgb.save(buf, format='PNG')
                return buf.getvalue()
            except Exception:
                return img_bytes
        except Exception:
            return img_bytes  # 预处理失败则原图传入

    @staticmethod
    def _detect_pdf_type(
        file_path: str,
        valid_pages: List[int],
        sample_pages: int = 5,
    ) -> str:
        """
        检测 PDF 类型：editable / scanned / mixed

        策略（双维度检测）：
        1. 文本长度：是否有足够多的嵌入式文本
        2. 文本质量：文本是否包含有意义的中文句子（而非纯数字/乱码）
        3. 页面图片：页面是否包含大尺寸嵌入图片（扫描件特征）

        - 全部无文本或文本质量极低 → scanned
        - 部分有文本 → mixed
        - 全部有高质量文本 → editable
        """
        try:
            import fitz
            doc = fitz.open(file_path)
        except ImportError:
            return "editable"  # 回退

        sample = valid_pages[:min(sample_pages, len(valid_pages))]
        if not sample:
            doc.close()
            return "editable"

        editable_count = 0
        scanned_count = 0

        for page_idx in sample:
            page = doc[page_idx]
            text = page.get_text().strip()
            text_len = len(text)
            quality = DocumentParser._text_quality_score(text)

            # 检查页面是否有嵌入图片（扫描件特征：整页图片）
            images = page.get_images(full=True)
            has_large_images = False
            if images:
                page_rect = page.rect
                page_area = page_rect.width * page_rect.height
                for img_info in images:
                    # 图片尺寸接近页面尺寸 → 扫描页
                    try:
                        img_rect = page.get_image_bbox(img_info)
                        if img_rect:
                            img_area = img_rect.width * img_rect.height
                            if img_area > page_area * 0.5:
                                has_large_images = True
                                break
                    except Exception:
                        pass

            # 判断逻辑：
            # 1. 文本质量高 → editable
            # 2. 文本极短 → scanned
            # 3. 文本长度够但质量差（纯数字/乱码）+ 有大幅图片 → scanned
            # 4. 文本长度够但质量差 + 无大幅图片 → 仍算 ambiguous，暂归 mixed
            if text_len > 100 and quality >= 0.15:
                editable_count += 1
            elif text_len < 10:
                scanned_count += 1
            elif has_large_images and quality < 0.15:
                # 有文本但质量差 + 大幅图片 → 扫描件
                scanned_count += 1
            elif quality < 0.08:
                # 文本质量极差 → 扫描件
                scanned_count += 1
            else:
                # 有少量文本，质量中等 → mixed
                # 不计数，算作模糊地带
                pass

        doc.close()

        total = len(sample)
        if total == 0:
            return "editable"

        scanned_ratio = scanned_count / total
        editable_ratio = editable_count / total

        if scanned_ratio >= 0.8:
            return "scanned"
        elif scanned_ratio >= 0.3 or (editable_ratio < 0.5 and scanned_ratio > 0):
            return "mixed"
        else:
            return "editable"

    @staticmethod
    def _parse_editable_pdf(
        file_path: str,
        valid_pages: List[int],
    ) -> str:
        """
        解析可编辑 PDF：主选 pdfplumber，备选 PyMuPDF，兜底 OCR

        关键改进：即使 _detect_pdf_type 判定为 editable，
        如果 pdfplumber/PyMuPDF 提取结果质量太差（纯数字、乱码），
        说明可能是扫描件 PDF 的嵌入垃圾文本层，应回退到 PaddleOCR。
        """
        best_text = ""
        best_source = ""

        # 主选：pdfplumber
        try:
            import pdfplumber
            text = DocumentParser._extract_with_pdfplumber(file_path, valid_pages)
            if text and len(text.strip()) > 100:
                quality = DocumentParser._text_quality_score(text)
                if quality >= 0.15:
                    logger.info(f"pdfplumber 解析成功: {len(text)} 字符, 质量={quality:.2f}")
                    return text
                elif len(text) > len(best_text):
                    best_text = text
                    best_source = "pdfplumber"
        except Exception as e:
            logger.warning(f"pdfplumber 解析失败，尝试 PyMuPDF: {str(e)}")

        # 备选：PyMuPDF
        try:
            text = DocumentParser._extract_with_fitz(file_path, valid_pages)
            if text and len(text.strip()) > 100:
                quality = DocumentParser._text_quality_score(text)
                if quality >= 0.15:
                    logger.info(f"PyMuPDF 解析成功: {len(text)} 字符, 质量={quality:.2f}")
                    return text
                elif len(text) > len(best_text):
                    best_text = text
                    best_source = "PyMuPDF"
        except Exception as e:
            logger.error(f"PyMuPDF 解析也失败: {str(e)}")

        # 可编辑 PDF 不应回退到 OCR：原生提取的文本即使格式不佳，
        # 也比 OCR 结果准确得多。OCR 仅适用于扫描件 PDF。
        if best_text and len(best_text) > 100:
            quality = DocumentParser._text_quality_score(best_text)
            logger.warning(
                f"{best_source} 提取质量偏低 (质量={quality:.2f})，但仍返回原生提取结果"
            )
            return best_text

        # 最终兜底：返回已有文本（即使质量差）
        return best_text if best_text else ""

    @staticmethod
    def _extract_with_pdfplumber(
        file_path: str,
        valid_pages: List[int],
    ) -> str:
        """使用 pdfplumber 提取 PDF 文本"""
        import pdfplumber

        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_idx in valid_pages:
                if page_idx >= len(pdf.pages):
                    continue
                page = pdf.pages[page_idx]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    # 基础行级清洗
                    lines = page_text.split('\n')
                    cleaned = []
                    for line in lines:
                        stripped = line.strip()
                        if not stripped:
                            continue
                        # 去除控制字符
                        stripped = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', stripped)
                        if not stripped.strip():
                            continue
                        cleaned.append(stripped)
                    if cleaned:
                        text_parts.append('\n'.join(cleaned))

        return '\n\n'.join(text_parts)

    @staticmethod
    def _extract_with_fitz(
        file_path: str,
        valid_pages: List[int],
    ) -> str:
        """使用 PyMuPDF (fitz) 提取 PDF 文本"""
        import fitz

        text_parts = []
        doc = fitz.open(file_path)
        try:
            for page_idx in valid_pages:
                if page_idx >= len(doc):
                    continue
                page = doc[page_idx]
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
        finally:
            doc.close()

        return '\n\n'.join(text_parts)

    @staticmethod
    def _parse_scanned_pdf(
        file_path: str,
        valid_pages: List[int],
        fitz_doc=None,
    ) -> str:
        """
        解析扫描版 PDF：使用 PaddleOCR

        低质量页面（识别率 < 60%）直接拦截。
        """
        from app.services.ocr_service import ocr_service

        if not ocr_service.is_available:
            logger.warning("PaddleOCR 不可用，尝试 PyMuPDF 提取")
            try:
                return DocumentParser._extract_with_fitz(file_path, valid_pages)
            except Exception:
                return ""

        doc = None
        try:
            if fitz_doc is not None:
                doc = fitz_doc
            else:
                import fitz
                doc = fitz.open(file_path)
        except ImportError:
            logger.error("PyMuPDF 不可用，无法渲染扫描 PDF")
            return ""

        text_parts = []
        rejected_count = 0

        try:
            for page_idx in valid_pages:
                if page_idx >= len(doc):
                    continue
                page = doc[page_idx]

                # 渲染页面为图片（150 DPI 平衡速度与质量）
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")

                # 图像预处理：增强对比度 + 自适应二值化
                img_bytes = DocumentParser._preprocess_ocr_image(img_bytes)

                # OCR 识别（带版面分析 + 置信度校验）
                text, confidence, layout_info = ocr_service.ocr_page_with_layout(
                    img_bytes,
                    min_confidence=DocumentParser.OCR_MIN_CONFIDENCE,
                )

                if layout_info.get("rejected"):
                    logger.warning(
                        f"扫描 PDF 第 {page_idx+1} 页被拦截: "
                        f"识别率 {confidence:.2%} < {DocumentParser.OCR_MIN_CONFIDENCE}"
                    )
                    rejected_count += 1
                    continue

                if text and len(text.strip()) >= DocumentParser.OCR_TEXT_MIN_LENGTH:
                    text_parts.append(text.strip())
                elif text:
                    logger.warning(f"扫描 PDF 第 {page_idx+1} 页内容过短 ({len(text)} 字符)，跳过")

        finally:
            # 只有自己打开的 doc 才关闭，外部传入的不关
            if fitz_doc is None and doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass

        if rejected_count > 0:
            logger.warning(f"扫描 PDF 共拦截 {rejected_count} 个低质量页面")

        return '\n\n'.join(text_parts)

    @staticmethod
    def _parse_mixed_pdf(
        file_path: str,
        valid_pages: List[int],
        fitz_doc=None,
    ) -> str:
        """
        解析混合 PDF（部分可编辑 + 部分扫描件）

        策略（正文用原生提取 + 图片区域单独 OCR）：
        - 先快速采样前 3 页，若全为扫描件 → 直接走 _parse_scanned_pdf（避免逐页开销）
        - 每页检测原生文本质量
        - 高质量文本 → 直接用原生提取（pdfplumber/PyMuPDF）
        - 页面内的独立图片 → 单独裁剪做 OCR，插入到对应位置
        - 原生文本质量差 + 无嵌入图片 → 整页渲染为图片走 OCR（兜底）
        """
        doc = None
        try:
            if fitz_doc is not None:
                doc = fitz_doc
            else:
                import fitz
                doc = fitz.open(file_path)
        except ImportError:
            return DocumentParser._parse_editable_pdf(file_path, valid_pages)

        from app.services.ocr_service import ocr_service

        # ---- 快速重分类：前 3 页全为扫描件 → 直接走扫描路径，避免逐页开销 ----
        sample_n = min(3, len(valid_pages))
        if sample_n > 0 and ocr_service.is_available:
            all_scanned = True
            for i in range(sample_n):
                page_idx = valid_pages[i]
                if page_idx < len(doc):
                    text = doc[page_idx].get_text().strip()
                    quality = DocumentParser._text_quality_score(text)
                    if len(text) > 50 and quality >= 0.15:
                        all_scanned = False
                        break
            if all_scanned:
                logger.info("混合 PDF 快速重分类：前 %d 页均为扫描件，切换为全量 OCR", sample_n)
                return DocumentParser._parse_scanned_pdf(file_path, valid_pages, fitz_doc=doc)

        text_parts = []

        try:
            for page_idx in valid_pages:
                if page_idx >= len(doc):
                    continue
                page = doc[page_idx]

                # 1. 提取原生文本并检测质量
                native_text = page.get_text().strip()
                native_quality = DocumentParser._text_quality_score(native_text)

                page_has_content = False

                # 2. 正文：原生文本质量高 → 直接使用
                if native_text and len(native_text) > 50 and native_quality >= 0.15:
                    text_parts.append(native_text)
                    page_has_content = True

                # 3. 图片区域：跳过嵌入图片 OCR（mixed PDF 中嵌入图片多为照片/背景，
                # 文字内容已在原生文本中；若原生文本不足，由第4步整页渲染 OCR 兜底）
                image_texts = []
                if not page_has_content and ocr_service.is_available:
                    # 仅当原生文本不可用时，才尝试 OCR 嵌入图片
                    image_list = page.get_images(full=True)
                    for img_info in image_list:
                        xref = img_info[0]
                        try:
                            base_image = doc.extract_image(xref)
                            img_bytes = base_image["image"]

                            # 跳过太小的图片（可能是图标/装饰）
                            if len(img_bytes) < 5000:
                                continue

                            # 预检：跳过低对比度图片（照片/背景等非文字内容）
                            if not DocumentParser._is_text_likely_image(img_bytes):
                                continue

                            # OCR 图片
                            img_text, img_conf, _ = ocr_service.ocr_page_with_layout(
                                img_bytes,
                                min_confidence=DocumentParser.OCR_MIN_CONFIDENCE,
                            )
                            if img_text and len(img_text.strip()) > 10:
                                image_texts.append(img_text.strip())

                        except Exception as e:
                            logger.debug(f"图片 OCR 失败: {str(e)}")

                if image_texts:
                    text_parts.append('\n'.join(image_texts))
                    page_has_content = True

                # 4. 兜底：原生文本质量差 + 无嵌入图片内容 → 整页 OCR
                if not page_has_content and ocr_service.is_available:
                    try:
                        pix = page.get_pixmap(dpi=150)
                        img_bytes = pix.tobytes("png")

                        # 图像预处理：增强对比度 + 自适应二值化
                        img_bytes = DocumentParser._preprocess_ocr_image(img_bytes)

                        ocr_text, ocr_conf, ocr_info = ocr_service.ocr_page_with_layout(
                            img_bytes,
                            min_confidence=DocumentParser.OCR_MIN_CONFIDENCE,
                        )
                        if not ocr_info.get("rejected") and ocr_text and len(ocr_text.strip()) > 30:
                            text_parts.append(ocr_text.strip())
                            logger.info(
                                f"混合 PDF 第 {page_idx+1} 页原生文本质量差 (quality={native_quality:.2f})，"
                                f"已回退到整页 OCR"
                            )
                    except Exception as e:
                        logger.debug(f"混合 PDF 整页 OCR 失败: {str(e)}")

        finally:
            # 只有自己打开的 doc 才关闭，外部传入的不关
            if fitz_doc is None and doc is not None:
                try:
                    doc.close()
                except Exception:
                    pass

        return '\n\n'.join(text_parts)

    # ==================== DOCX 解析 ====================

    @staticmethod
    def _load_docx_image_map(file_path: str) -> Dict[str, bytes]:
        """
        从 DOCX（ZIP）中加载 rId → 图片字节 的映射。

        DOCX 内的图片存储在 word/media/ 下，通过
        word/_rels/document.xml.rels 中的 Relationship 关联到 rId。
        """
        import zipfile
        from xml.etree import ElementTree

        image_map: Dict[str, bytes] = {}
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 解析关系文件，建立 rId → 图片路径的映射
                rels_path = 'word/_rels/document.xml.rels'
                if rels_path not in zf.namelist():
                    return image_map

                rels_xml = zf.read(rels_path)
                rels_tree = ElementTree.fromstring(rels_xml)
                ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'

                rid_to_path = {}
                for rel in rels_tree:
                    rid = rel.attrib.get('Id', '')
                    target = rel.attrib.get('Target', '')
                    rel_type = rel.attrib.get('Type', '')
                    if 'image' in rel_type.lower() or target.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')):
                        # target 可能是相对路径，如 "media/image1.png"
                        full_path = f"word/{target}" if not target.startswith('word/') else target
                        rid_to_path[rid] = full_path

                # 读取图片字节
                for rid, img_path in rid_to_path.items():
                    if img_path in zf.namelist():
                        image_map[rid] = zf.read(img_path)

        except Exception as e:
            logger.warning(f"加载 DOCX 图片映射失败: {str(e)}")

        return image_map

    @staticmethod
    def _ocr_docx_images(
        image_map: Dict[str, bytes],
        ocr_service,
    ) -> Dict[str, str]:
        """
        对 DOCX 中提取的图片逐个 OCR，返回 rId → OCR 文本 的映射。
        只 OCR 大小 >= 5KB 的图片（跳过图标/装饰）。
        """
        ocr_map: Dict[str, str] = {}
        if not ocr_service.is_available:
            return ocr_map

        for rid, img_bytes in image_map.items():
            if len(img_bytes) < 5000:
                continue
            try:
                img_text, img_conf, ocr_info = ocr_service.ocr_page_with_layout(
                    img_bytes,
                    min_confidence=DocumentParser.OCR_MIN_CONFIDENCE,
                )
                if not ocr_info.get("rejected") and img_text and len(img_text.strip()) > 10:
                    ocr_map[rid] = img_text.strip()
            except Exception as e:
                logger.debug(f"DOCX 图片 OCR 失败 (rId={rid}): {str(e)}")

        return ocr_map

    @staticmethod
    def _get_drawing_rids(paragraph_element) -> List[str]:
        """
        从段落 XML 元素中提取所有 w:drawing 引用的图片 rId。

        路径：w:drawing > wp:inline > a:graphic > a:graphicData > pic:pic > pic:blipFill > a:blip[@r:embed]
        """
        rids = []
        nsmap = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        for blip in paragraph_element.findall('.//a:blip', nsmap):
            rid = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rid:
                rids.append(rid)
        return rids

    @staticmethod
    def _parse_docx(
        file_path: str,
        preprocess: bool = True,
    ) -> str:
        """
        解析 DOCX 文件（主选 python-docx）

        策略：正文用原生提取 + 图片区域单独 OCR
        - 段落文本 → 原生提取
        - 段落中的嵌入图片 → 单独 OCR，插入到段落文本之后
        - 表格 → 提取结构化文本
        """
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("python-docx 不可用")
            return ""

        from app.services.ocr_service import ocr_service

        doc = DocxDocument(file_path)

        # 预处理：获取有效段落索引（基于 doc.paragraphs 顺序）
        valid_para_indices = None
        if preprocess:
            valid_para_indices = document_preprocessor.preprocess_docx_paragraphs(file_path)

        # 加载图片映射并批量 OCR（提前做，避免重复解析 ZIP）
        image_map = DocumentParser._load_docx_image_map(file_path)
        ocr_map = DocumentParser._ocr_docx_images(image_map, ocr_service) if image_map else {}

        text_parts = []
        para_idx = 0  # doc.paragraphs 中的索引计数器

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # 当前段落索引
                if valid_para_indices is not None and para_idx not in valid_para_indices:
                    para_idx += 1
                    continue

                # 直接从 doc.paragraphs 按索引获取，避免 O(n²) 查找
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para_text = para.text.strip()

                    # 正文：原生提取段落文本
                    if para_text:
                        text_parts.append(para_text)

                    # 图片区域：提取段落中的嵌入图片并 OCR
                    if ocr_map:
                        drawing_rids = DocumentParser._get_drawing_rids(element)
                        ocr_texts = []
                        for rid in drawing_rids:
                            if rid in ocr_map:
                                ocr_texts.append(ocr_map[rid])
                        if ocr_texts:
                            text_parts.append('\n'.join(ocr_texts))

                para_idx += 1

            elif tag == 'tbl':
                # 表格不参与段落索引计数，直接提取
                table_text = DocumentParser._extract_table_text(element, doc)
                if table_text:
                    text_parts.append(table_text)

        return '\n\n'.join(text_parts)

    @staticmethod
    def _extract_table_text(table_element, doc) -> str:
        """提取表格文本并保持结构"""
        try:
            from docx.oxml.ns import qn
        except ImportError:
            return ""

        rows = []
        for tr in table_element.findall('.//' + qn('w:tr')):
            cells = []
            for tc in tr.findall('.//' + qn('w:tc')):
                cell_text = ''.join(node.text or '' for node in tc.iter() if node.text)
                cells.append(cell_text.strip())
            if cells:
                rows.append(' | '.join(cells))

        if not rows:
            return ""

        if len(rows) > 1:
            header_parts = ['---'] * len(rows[0].split(' | '))
            rows.insert(1, ' | '.join(header_parts))

        return '\n'.join(rows)

    # ==================== TXT / MD 解析 ====================

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """解析 TXT 文件（多编码尝试）"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        return ""

    @staticmethod
    def _parse_md(file_path: str) -> str:
        """解析 MD 文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception:
                return ""

    # ==================== 流式解析 ====================

    @staticmethod
    def parse_stream(
        file_path: str,
        chunk_pages: int = 5,
        preprocess: bool = True,
    ) -> Generator[str, None, None]:
        """流式解析文档"""
        validation = DocumentParser.validate_file(file_path)
        if not validation["valid"]:
            raise ValueError(f"文件校验失败: {validation['error']}")

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            yield from DocumentParser._parse_pdf_stream(file_path, chunk_pages, preprocess)
        elif ext in [".docx", ".doc"]:
            yield from DocumentParser._parse_docx_stream(file_path, chunk_pages, preprocess)
        else:
            yield DocumentParser.parse(file_path, preprocess=preprocess, apply_coarse_clean=True)

    @staticmethod
    def _parse_pdf_stream(
        file_path: str,
        chunk_pages: int,
        preprocess: bool,
    ) -> Generator[str, None, None]:
        """流式解析 PDF"""
        valid_pages = document_preprocessor.preprocess_pdf_pages(file_path) if preprocess else None

        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                batch = []
                for i, page in enumerate(pdf.pages):
                    if valid_pages and i not in valid_pages:
                        continue
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        batch.append(page_text.strip())
                    if len(batch) >= chunk_pages:
                        yield '\n\n'.join(batch)
                        batch = []
                if batch:
                    yield '\n\n'.join(batch)
        except Exception:
            try:
                import fitz
                doc = fitz.open(file_path)
                batch = []
                for i in (valid_pages or range(len(doc))):
                    if i >= len(doc):
                        continue
                    page_text = doc[i].get_text()
                    if page_text and page_text.strip():
                        batch.append(page_text.strip())
                    if len(batch) >= chunk_pages:
                        yield '\n\n'.join(batch)
                        batch = []
                if batch:
                    yield '\n\n'.join(batch)
                doc.close()
            except Exception:
                return

    @staticmethod
    def _parse_docx_stream(
        file_path: str,
        chunk_pages: int,
        preprocess: bool,
    ) -> Generator[str, None, None]:
        """
        流式解析 DOCX（正文用原生提取 + 图片区域单独 OCR）
        """
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError:
            return

        from app.services.ocr_service import ocr_service

        doc = DocxDocument(file_path)
        valid_para_indices = document_preprocessor.preprocess_docx_paragraphs(file_path) if preprocess else None

        # 加载图片映射并批量 OCR
        image_map = DocumentParser._load_docx_image_map(file_path)
        ocr_map = DocumentParser._ocr_docx_images(image_map, ocr_service) if image_map else {}

        batch = []
        count = 0
        para_idx = 0

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                if valid_para_indices is not None and para_idx not in valid_para_indices:
                    para_idx += 1
                    continue

                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para_text = para.text.strip()

                    # 正文：原生提取段落文本
                    if para_text:
                        batch.append(para_text)
                        count += 1

                    # 图片区域：提取段落中的嵌入图片并 OCR
                    if ocr_map:
                        drawing_rids = DocumentParser._get_drawing_rids(element)
                        ocr_texts = []
                        for rid in drawing_rids:
                            if rid in ocr_map:
                                ocr_texts.append(ocr_map[rid])
                        if ocr_texts:
                            batch.append('\n'.join(ocr_texts))
                            count += 1

                para_idx += 1

            elif tag == 'tbl':
                table_text = DocumentParser._extract_table_text(element, doc)
                if table_text:
                    batch.append(table_text)
                    count += 1

            if count >= chunk_pages:
                yield '\n\n'.join(batch)
                batch = []
                count = 0

        if batch:
            yield '\n\n'.join(batch)

    # ==================== 大文件检测 ====================

    @staticmethod
    def is_large_file(file_path: str) -> bool:
        """检查文件是否需要拆分（>50MB）"""
        try:
            return os.path.getsize(file_path) > DocumentParser.LARGE_FILE_THRESHOLD
        except Exception:
            return False

    @staticmethod
    def get_file_size(file_path: str) -> int:
        """获取文件大小"""
        try:
            return os.path.getsize(file_path)
        except Exception:
            return 0