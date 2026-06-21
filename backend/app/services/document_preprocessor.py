"""
文档预处理模块
==============
在上传后、解析前，自动剔除文档中的无效/重复内容：
- 封面页、扉页
- 目录页
- 版权声明页
- 参考文献页
- 附录页
- 纯空白页
- 页眉页脚、页码、水印

支持 PDF 和 DOCX 格式。
"""

import logging
import os
import re
from typing import List, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


# 封面关键词（出现这些词且页面内容很少 → 封面/扉页）
COVER_KEYWORDS = [
    '封面', '扉页', '书名', '编著', '主编', '出版社',
    '出版日期', '版次', '印次', 'ISBN', 'CIP',
]

# 目录特征
TOC_HEADER_KEYWORDS = ['目录', '目次', 'Table of Contents', 'CONTENTS']

# 版权声明关键词
COPYRIGHT_KEYWORDS = [
    '版权所有', '版权声明', '翻印必究', '侵权必究',
    'Copyright', 'All Rights Reserved', '未经许可',
    '不得以任何形式', '最终解释权',
]

# 参考文献关键词
REFERENCE_KEYWORDS = [
    '参考文献', 'References', '参考资料',
    '引用文献', '参考书目', 'Bibliography',
]

# 附录关键词
APPENDIX_KEYWORDS = [
    '附录', 'Appendix', '附件', '附表',
    '附则', '补充说明',
]

# 常见页眉页脚关键词（跨页重复出现）
HEADER_FOOTER_COMMON = [
    '机密', '内部资料', '版权所有', '未经许可',
    '文档编号', '版本号', '修订日期', '生效日期',
    'Confidential', 'Internal', 'Draft',
]


class DocumentPreprocessor:
    """
    文档预处理器
    
    在解析前识别并跳过无效页面/段落。
    """

    def __init__(self):
        self._seen_header_footer: set = set()

    # ==================== PDF 预处理 ====================

    def preprocess_pdf_pages(
        self,
        pdf_path: str,
        remove_cover: bool = True,
        remove_toc: bool = True,
        remove_references: bool = True,
        remove_appendix: bool = True,
        remove_blank: bool = True,
    ) -> List[int]:
        """
        预处理 PDF，返回有效页面的页码列表（0-based）。
        
        Args:
            pdf_path: PDF 文件路径
            remove_cover: 是否移除封面/扉页
            remove_toc: 是否移除目录
            remove_references: 是否移除参考文献
            remove_appendix: 是否移除附录
            remove_blank: 是否移除空白页
        
        Returns:
            有效页面索引列表
        """
        valid_pages = []
        self._seen_header_footer = set()

        try:
            import fitz
            doc = fitz.open(pdf_path)
        except ImportError:
            logger.warning("PyMuPDF 不可用，跳过 PDF 预处理")
            try:
                import pdfplumber
                with pdfplumber.open(pdf_path) as pdf:
                    return list(range(len(pdf.pages)))
            except Exception:
                return []

        total_pages = len(doc)
        if total_pages == 0:
            doc.close()
            return []

        # 阶段1：收集所有页面的文本特征
        page_texts = []
        for i in range(total_pages):
            page = doc[i]
            text = page.get_text()
            page_texts.append(text)

        # 阶段2：识别各区域
        cover_pages = set()
        toc_pages = set()
        reference_pages = set()
        appendix_pages = set()
        blank_pages = set()

        for i, text in enumerate(page_texts):
            text_stripped = text.strip()

            # 空白页检测
            if remove_blank and (not text_stripped or len(text_stripped) < 20):
                blank_pages.add(i)
                continue

            text_lower = text_stripped.lower()

            # 封面/扉页检测（前几页 + 关键词）
            if remove_cover and i < max(5, total_pages * 0.1):
                cover_score = sum(1 for kw in COVER_KEYWORDS if kw.lower() in text_lower)
                if cover_score >= 2 or (cover_score >= 1 and len(text_stripped) < 200):
                    cover_pages.add(i)
                    continue

            # 目录检测
            if remove_toc:
                is_toc_header = any(kw.lower() in text_lower for kw in TOC_HEADER_KEYWORDS)
                if is_toc_header:
                    toc_pages.add(i)
                    # 标记目录页后的连续页面（直到遇到正文）
                    for j in range(i + 1, min(i + 15, total_pages)):
                        next_text = page_texts[j].strip()
                        if any(kw.lower() in next_text.lower() for kw in REFERENCE_KEYWORDS + APPENDIX_KEYWORDS):
                            break
                        if len(next_text) > 100 and not self._looks_like_toc(next_text):
                            break
                        toc_pages.add(j)
                    continue

            # 参考文献检测
            if remove_references:
                if any(kw.lower() in text_lower for kw in REFERENCE_KEYWORDS):
                    reference_pages.add(i)
                    for j in range(i + 1, total_pages):
                        ref_text = page_texts[j].strip()
                        if any(kw.lower() in ref_text.lower() for kw in APPENDIX_KEYWORDS):
                            break
                        if len(ref_text) > 50 and not self._looks_like_reference(ref_text):
                            break
                        reference_pages.add(j)
                    continue

            # 附录检测
            if remove_appendix:
                if any(kw.lower() in text_lower for kw in APPENDIX_KEYWORDS):
                    # 只有关键词作为标题出现时才触发（页面内容短，非正文中偶然提及）
                    is_appendix_heading = len(text_stripped) < 300
                    if is_appendix_heading:
                        appendix_pages.add(i)
                        for j in range(i + 1, total_pages):
                            next_text = page_texts[j].strip()
                            next_lower = next_text.lower()
                            # 停止条件1：遇到参考文献
                            if any(kw.lower() in next_lower for kw in REFERENCE_KEYWORDS):
                                break
                            # 停止条件2：页面内容较长且不含附录关键词，可能已回到正文
                            if len(next_text) > 500 and not any(
                                kw.lower() in next_lower for kw in APPENDIX_KEYWORDS
                            ):
                                break
                            appendix_pages.add(j)
                    continue

            # 版权声明检测（整页都是版权声明）
            if any(kw.lower() in text_lower for kw in COPYRIGHT_KEYWORDS) and len(text_stripped) < 300:
                continue

        # 阶段3：构建有效页面列表
        all_removed = cover_pages | toc_pages | reference_pages | appendix_pages | blank_pages

        for i in range(total_pages):
            if i not in all_removed:
                valid_pages.append(i)

        doc.close()

        removed_count = len(all_removed)
        if removed_count > 0:
            logger.info(
                f"PDF 预处理: 共 {total_pages} 页, "
                f"移除 {removed_count} 页 "
                f"(封面{len(cover_pages)} 目录{len(toc_pages)} "
                f"参考文献{len(reference_pages)} 附录{len(appendix_pages)} "
                f"空白{len(blank_pages)}), "
                f"保留 {len(valid_pages)} 页"
            )

        return valid_pages

    @staticmethod
    def _looks_like_toc(text: str) -> bool:
        """判断文本是否像目录内容"""
        lines = text.strip().split('\n')
        toc_line_count = 0
        for line in lines[:20]:
            stripped = line.strip()
            if not stripped:
                continue
            # 目录特征：引导点线 + 页码
            if re.search(r'[.·•]{3,}', stripped):
                toc_line_count += 1
            elif re.search(r'\.{2,}\s*\d+$', stripped):
                toc_line_count += 1
            elif re.match(r'^.{2,30}\s+\d{1,4}$', stripped):
                toc_line_count += 1
        return toc_line_count >= 3

    @staticmethod
    def _looks_like_reference(text: str) -> bool:
        """判断文本是否像参考文献"""
        lines = text.strip().split('\n')
        ref_count = 0
        for line in lines[:15]:
            stripped = line.strip()
            if not stripped:
                continue
            # 参考文献特征：[1], [2] 或 1. 开头
            if re.match(r'^\[\d+\]', stripped):
                ref_count += 1
            elif re.match(r'^\d+\.\s', stripped) and len(stripped) > 20:
                ref_count += 1
        return ref_count >= 3

    # ==================== DOCX 预处理 ====================

    def preprocess_docx_paragraphs(
        self,
        docx_path: str,
        remove_cover: bool = True,
        remove_toc: bool = True,
        remove_references: bool = True,
        remove_appendix: bool = True,
    ) -> List[int]:
        """
        预处理 DOCX，返回有效段落的索引列表。
        
        Args:
            docx_path: DOCX 文件路径
            remove_cover: 是否移除封面相关段落
            remove_toc: 是否移除目录
            remove_references: 是否移除参考文献
            remove_appendix: 是否移除附录
        
        Returns:
            有效段落索引列表
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx 不可用，跳过 DOCX 预处理")
            return []

        doc = DocxDocument(docx_path)
        total_paras = len(doc.paragraphs)
        if total_paras == 0:
            return []

        valid_paras = []
        in_toc = False
        in_ref = False
        in_appendix = False

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            text_lower = text.lower()

            if not text:
                continue

            # 封面检测（前20个段落 + 关键词）
            if remove_cover and i < 20:
                if any(kw.lower() in text_lower for kw in COVER_KEYWORDS) and len(text) < 100:
                    continue

            # 目录检测
            if remove_toc:
                if any(kw.lower() in text_lower for kw in TOC_HEADER_KEYWORDS):
                    in_toc = True
                    continue
                if in_toc:
                    if self._looks_like_toc(text) or len(text) < 80:
                        continue
                    else:
                        in_toc = False

            # 参考文献检测
            if remove_references:
                if any(kw.lower() in text_lower for kw in REFERENCE_KEYWORDS):
                    in_ref = True
                    continue
                if in_ref:
                    if self._looks_like_reference(text):
                        continue
                    else:
                        in_ref = False

            # 附录检测
            if remove_appendix:
                if any(kw.lower() in text_lower for kw in APPENDIX_KEYWORDS):
                    # 只有关键词作为标题出现时才触发（段落短，非正文中偶然提及）
                    if len(text) < 300:
                        in_appendix = True
                        continue
                if in_appendix:
                    # 停止条件1：遇到参考文献
                    if remove_references and any(kw.lower() in text_lower for kw in REFERENCE_KEYWORDS):
                        in_appendix = False
                        in_ref = True
                        continue
                    # 停止条件2：段落较长且不含附录关键词，可能已回到正文
                    if len(text) > 500 and not any(kw.lower() in text_lower for kw in APPENDIX_KEYWORDS):
                        in_appendix = False
                    else:
                        continue

            # 版权声明
            if any(kw.lower() in text_lower for kw in COPYRIGHT_KEYWORDS) and len(text) < 200:
                continue

            valid_paras.append(i)

        removed = total_paras - len(valid_paras)
        if removed > 0:
            logger.info(f"DOCX 预处理: 共 {total_paras} 段, 移除 {removed} 段, 保留 {len(valid_paras)} 段")

        return valid_paras

    # ==================== 图片预处理 ====================

    def preprocess_image_page(
        self,
        image_bytes: bytes,
        return_early_exit: bool = True,
    ) -> Tuple[bool, str]:
        """
        判断图片页面是否需要处理
        
        Args:
            image_bytes: 图片字节
            return_early_exit: 是否提前退出（空白/低质量直接返回）
        
        Returns:
            (should_process, reason)
        """
        try:
            from PIL import Image
            import io
            import numpy as np

            img = Image.open(io.BytesIO(image_bytes))
            img_array = np.array(img.convert('L'))

            # 空白页检测：像素均值接近255
            mean_brightness = np.mean(img_array)
            if mean_brightness > 250:
                return False, "blank_page"

            # 低质量模糊检测：拉普拉斯方差
            if return_early_exit:
                laplacian_var = np.var(np.gradient(np.gradient(img_array.astype(float))))
                if laplacian_var < 10:
                    return False, "blurry"

            return True, "ok"

        except Exception as e:
            logger.warning(f"图片预处理异常: {str(e)}")
            return True, "ok"


# 全局单例
document_preprocessor = DocumentPreprocessor()